"""
Document processing service.
Handles file upload, validation, content extraction, and processing orchestration.
"""
import os
import hashlib
import mimetypes
import psutil
import resource
from typing import Optional, Tuple, Dict, Any, List
from pathlib import Path
import asyncio
from datetime import datetime
from contextlib import asynccontextmanager
from dataclasses import dataclass, field

from app.core.config import settings
from app.core.logging import get_logger, get_utc_datetime
from app.core.database import db_manager
from app.core.exceptions import (
    DocumentProcessingError,
    FileTooLargeError,
    UnsupportedFileTypeError,
    ValidationError
)
from app.models.document import Document, DocumentChunk
from app.schemas.document import DocumentCreate, ProcessingStatus
from app.services.metrics import metrics_service
from sqlalchemy import select, update, delete
from sqlalchemy.ext.asyncio import AsyncSession

logger = get_logger(__name__)


@dataclass
class ProcessingTransaction:
    """
    Tracks processing operations for atomic rollback.
    """
    document_id: Optional[str] = None
    created_files: List[str] = field(default_factory=list)
    created_chunks: List[str] = field(default_factory=list)
    created_embeddings: List[str] = field(default_factory=list)
    temp_files: List[str] = field(default_factory=list)
    session: Optional[AsyncSession] = None
    is_committed: bool = False
    rollback_errors: List[str] = field(default_factory=list)
    
    def add_file(self, filepath: str, is_temp: bool = False) -> None:
        """Track a created file."""
        if is_temp:
            self.temp_files.append(filepath)
        else:
            self.created_files.append(filepath)
    
    def add_chunk(self, chunk_id: str) -> None:
        """Track a created chunk."""
        self.created_chunks.append(chunk_id)
    
    def add_embedding(self, embedding_id: str) -> None:
        """Track a created embedding."""
        self.created_embeddings.append(embedding_id)
    
    async def rollback(self) -> bool:
        """
        Rollback all operations in reverse order.
        Returns True if rollback was successful, False if there were errors.
        """
        if self.is_committed:
            logger.warning("attempted_rollback_on_committed_transaction")
            return False
        
        logger.info(
            "starting_transaction_rollback",
            document_id=self.document_id,
            files=len(self.created_files),
            chunks=len(self.created_chunks),
            embeddings=len(self.created_embeddings),
            temp_files=len(self.temp_files)
        )
        
        # 1. Delete embeddings (if vector store supports it)
        for embedding_id in reversed(self.created_embeddings):
            try:
                # TODO: Implement vector store deletion when available
                logger.debug("would_delete_embedding", embedding_id=embedding_id)
            except Exception as e:
                self.rollback_errors.append(f"Failed to delete embedding {embedding_id}: {str(e)}")
        
        # 2. Delete chunks from database
        if self.session and self.created_chunks:
            try:
                await self.session.execute(
                    delete(DocumentChunk).where(
                        DocumentChunk.id.in_(self.created_chunks)
                    )
                )
                logger.info("chunks_deleted", count=len(self.created_chunks))
            except Exception as e:
                self.rollback_errors.append(f"Failed to delete chunks: {str(e)}")
        
        # 3. Delete document record
        if self.session and self.document_id:
            try:
                await self.session.execute(
                    delete(Document).where(Document.id == self.document_id)
                )
                logger.info("document_deleted", document_id=self.document_id)
            except Exception as e:
                self.rollback_errors.append(f"Failed to delete document: {str(e)}")
        
        # 4. Delete created files
        for filepath in reversed(self.created_files):
            try:
                if os.path.exists(filepath):
                    os.remove(filepath)
                    logger.debug("file_deleted", filepath=filepath)
            except Exception as e:
                self.rollback_errors.append(f"Failed to delete file {filepath}: {str(e)}")
        
        # 5. Clean up temp files
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
                    logger.debug("temp_file_deleted", filepath=temp_file)
            except Exception as e:
                self.rollback_errors.append(f"Failed to delete temp file {temp_file}: {str(e)}")
        
        # 6. Rollback database transaction if needed
        if self.session:
            try:
                await self.session.rollback()
            except Exception as e:
                self.rollback_errors.append(f"Database rollback failed: {str(e)}")
        
        success = len(self.rollback_errors) == 0
        
        logger.info(
            "transaction_rollback_completed",
            success=success,
            errors=len(self.rollback_errors),
            error_details=self.rollback_errors[:5] if self.rollback_errors else None
        )
        
        return success
    
    def commit(self) -> None:
        """Mark transaction as committed (no rollback allowed)."""
        self.is_committed = True
        logger.info(
            "transaction_committed",
            document_id=self.document_id,
            files=len(self.created_files),
            chunks=len(self.created_chunks)
        )


class DocumentProcessor:
    """Handles document processing operations with memory monitoring."""
    
    def __init__(self):
        self.supported_types = set(settings.ALLOWED_FILE_TYPES)
        self.max_size_bytes = settings.MAX_FILE_SIZE_MB * 1024 * 1024
        # Memory limits for processing (in MB)
        self.max_memory_mb = getattr(settings, 'MAX_PROCESSING_MEMORY_MB', 512)
        self.memory_check_interval = 100  # Check memory every N operations
    
    def get_memory_usage_mb(self) -> float:
        """Get current process memory usage in MB."""
        process = psutil.Process(os.getpid())
        return process.memory_info().rss / 1024 / 1024
    
    def check_memory_limit(self, operation: str = "processing") -> None:
        """Check if memory usage is within limits."""
        current_memory_mb = self.get_memory_usage_mb()
        
        if current_memory_mb > self.max_memory_mb:
            raise DocumentProcessingError(
                f"Memory limit exceeded during {operation}: "
                f"{current_memory_mb:.1f}MB > {self.max_memory_mb}MB limit"
            )
        
        # Warn if approaching limit (80%)
        if current_memory_mb > self.max_memory_mb * 0.8:
            logger.warning(
                "memory_usage_high",
                operation=operation,
                current_mb=current_memory_mb,
                limit_mb=self.max_memory_mb,
                percentage=(current_memory_mb / self.max_memory_mb) * 100
            )
    
    def estimate_required_memory(self, file_size_bytes: int, file_type: str) -> float:
        """Estimate memory required for processing a file."""
        # Base estimation: file needs to be loaded + overhead for processing
        # Different multipliers for different file types
        multipliers = {
            '.pdf': 3.0,   # PDF processing can use 3x file size
            '.docx': 2.5,  # DOCX processing uses ~2.5x
            '.txt': 1.5    # Text processing is most efficient
        }
        
        multiplier = multipliers.get(file_type, 2.0)
        return (file_size_bytes / 1024 / 1024) * multiplier
    
    def get_disk_usage(self, path: str = '/tmp') -> Dict[str, float]:
        """Get disk usage statistics for the given path."""
        try:
            usage = psutil.disk_usage(path)
            return {
                'total_gb': usage.total / (1024**3),
                'used_gb': usage.used / (1024**3),
                'free_gb': usage.free / (1024**3),
                'percent_used': usage.percent
            }
        except Exception as e:
            logger.error("disk_usage_check_failed", path=path, error=str(e))
            return {'free_gb': 0, 'percent_used': 100}
    
    def check_disk_space(self, required_bytes: int, path: str = '/tmp') -> None:
        """Check if sufficient disk space is available."""
        # Add buffer for safety (require 2x the file size)
        required_with_buffer = required_bytes * 2
        required_gb = required_with_buffer / (1024**3)
        
        disk_info = self.get_disk_usage(path)
        free_gb = disk_info.get('free_gb', 0)
        
        # Require at least 100MB free after operation
        min_free_gb = 0.1
        
        if free_gb < (required_gb + min_free_gb):
            raise DocumentProcessingError(
                f"Insufficient disk space. Required: {required_gb:.2f}GB, "
                f"Available: {free_gb:.2f}GB"
            )
        
        # Warn if disk usage is high (>90%)
        if disk_info.get('percent_used', 100) > 90:
            logger.warning(
                "disk_space_low",
                path=path,
                free_gb=free_gb,
                percent_used=disk_info.get('percent_used', 100)
            )
    
    def check_storage_permissions(self, path: str = '/tmp') -> None:
        """Check if we have write permissions to the storage path."""
        try:
            # Check if path exists
            if not os.path.exists(path):
                try:
                    os.makedirs(path, exist_ok=True)
                except Exception as e:
                    raise DocumentProcessingError(
                        f"Cannot create storage directory {path}: {str(e)}"
                    )
            
            # Check write permission
            test_file = os.path.join(path, f'.write_test_{os.getpid()}')
            try:
                with open(test_file, 'w') as f:
                    f.write('test')
                os.remove(test_file)
            except Exception as e:
                raise DocumentProcessingError(
                    f"No write permission for storage path {path}: {str(e)}"
                )
            
            # Check if path is actually a directory
            if not os.path.isdir(path):
                raise DocumentProcessingError(
                    f"Storage path {path} is not a directory"
                )
                
        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(
                f"Storage permission check failed: {str(e)}"
            )
    
    async def cleanup_partial_failure(
        self,
        document_id: Optional[str] = None,
        temp_files: Optional[List[str]] = None,
        session: Optional[AsyncSession] = None
    ) -> None:
        """
        Clean up resources after a partial failure.
        
        Args:
            document_id: Document ID to clean up from database
            temp_files: List of temporary files to remove
            session: Database session for cleanup operations
        """
        cleanup_errors = []
        
        # Clean up temporary files
        if temp_files:
            for temp_file in temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                        logger.info("temp_file_removed", file=temp_file)
                except Exception as e:
                    cleanup_errors.append(f"Failed to remove {temp_file}: {str(e)}")
        
        # Clean up database records
        if document_id and session:
            try:
                # Delete document chunks first (if any)
                await session.execute(
                    f"DELETE FROM document_chunks WHERE document_id = '{document_id}'"
                )
                
                # Delete document record
                await session.execute(
                    f"DELETE FROM documents WHERE id = '{document_id}'"
                )
                
                await session.commit()
                logger.info("document_cleanup_complete", document_id=document_id)
                
            except Exception as e:
                cleanup_errors.append(f"Database cleanup failed: {str(e)}")
                try:
                    await session.rollback()
                except:
                    pass
        
        # Log any cleanup errors
        if cleanup_errors:
            logger.error(
                "cleanup_errors_occurred",
                errors=cleanup_errors,
                document_id=document_id
            )
        
    async def validate_file(self, filename: str, content: bytes) -> None:
        """
        Validate uploaded file meets requirements with enhanced security checks.
        
        Args:
            filename: Original filename
            content: File content as bytes
            
        Raises:
            FileTooLargeError: If file exceeds size limit
            UnsupportedFileTypeError: If file type not supported
            ValidationError: If file fails security checks
        """
        # Check file size
        file_size_bytes = len(content)
        if file_size_bytes == 0:
            raise ValidationError(f"Empty file uploaded: {filename}")
            
        if file_size_bytes > self.max_size_bytes:
            file_size_mb = file_size_bytes / (1024 * 1024)
            raise FileTooLargeError(
                filename=filename,
                size_mb=file_size_mb,
                max_size_mb=settings.MAX_FILE_SIZE_MB
            )
        
        # Check file extension
        file_extension = Path(filename).suffix.lower()
        
        # Check memory before processing
        estimated_memory = self.estimate_required_memory(file_size_bytes, file_extension)
        current_memory = self.get_memory_usage_mb()
        
        if current_memory + estimated_memory > self.max_memory_mb:
            raise DocumentProcessingError(
                f"Insufficient memory to process file. Current: {current_memory:.1f}MB, "
                f"Required: {estimated_memory:.1f}MB, Limit: {self.max_memory_mb}MB"
            )
        
        # Check disk space before processing
        self.check_disk_space(file_size_bytes)
        
        if file_extension not in self.supported_types:
            raise UnsupportedFileTypeError(
                filename=filename,
                file_type=file_extension,
                supported_types=list(self.supported_types)
            )
        
        # MIME type verification
        mime_type = mimetypes.guess_type(filename)[0]
        expected_mimes = {
            '.pdf': ['application/pdf'],
            '.docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
            '.txt': ['text/plain', 'text/csv', 'text/tab-separated-values']
        }
        
        if file_extension in expected_mimes:
            valid_mimes = expected_mimes[file_extension]
            if mime_type and mime_type not in valid_mimes:
                raise ValidationError(
                    f"MIME type mismatch: expected {valid_mimes} but got {mime_type} for {filename}"
                )
        
        # Content-based validation (basic magic number checking)
        if file_extension == '.pdf' and not content.startswith(b'%PDF'):
            raise ValidationError(f"Invalid PDF file format: {filename}")
        elif file_extension == '.docx':
            # DOCX files are ZIP archives starting with PK
            if not content.startswith(b'PK'):
                raise ValidationError(f"Invalid DOCX file format: {filename}")
        
        # Check for malicious patterns (basic implementation)
        malicious_patterns = [
            b'<script',  # JavaScript in uploaded files
            b'<%',       # Server-side scripting
            b'<?php',    # PHP code
            b'\x00\x00\x00\x00',  # Null bytes that might indicate corruption
        ]
        
        for pattern in malicious_patterns:
            if pattern in content[:1024]:  # Check first 1KB
                raise ValidationError(f"Potentially malicious content detected in {filename}")
        
        logger.info(
            "file_validation_passed",
            filename=filename,
            size_bytes=file_size_bytes,
            file_type=file_extension,
            mime_type=mime_type
        )
    
    def calculate_file_hash(self, content: bytes) -> str:
        """Calculate SHA-256 hash of file content."""
        return hashlib.sha256(content).hexdigest()
    
    def generate_safe_filename(self, original_filename: str) -> str:
        """Generate a safe filename for storage with directory traversal prevention."""
        # Prevent directory traversal attacks
        if any(dangerous in original_filename for dangerous in ['..', '/', '\\', '\x00']):
            raise ValidationError(f"Filename contains dangerous characters: {original_filename}")
        
        # Extract just the filename component (no path)
        base_filename = os.path.basename(original_filename)
        
        # Remove unsafe characters and create unique filename
        safe_name = "".join(c for c in base_filename if c.isalnum() or c in "._-")
        if not safe_name:
            raise ValidationError(f"Filename contains no valid characters: {original_filename}")
            
        # REFACTORED: Using existing utility instead of direct datetime.utcnow()
        timestamp = get_utc_datetime().strftime("%Y%m%d_%H%M%S")
        name_part, ext_part = os.path.splitext(safe_name)
        
        # Ensure extension is safe
        if ext_part and ext_part not in self.supported_types:
            raise ValidationError(f"Invalid file extension: {ext_part}")
            
        return f"{timestamp}_{name_part[:50]}{ext_part}"
    
    async def extract_text_content(self, filename: str, content: bytes) -> Tuple[str, Optional[str]]:
        """
        Extract text content from file.
        
        Args:
            filename: Original filename
            content: File content as bytes
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        file_extension = Path(filename).suffix.lower()
        
        try:
            if file_extension == ".txt":
                return await self._extract_from_txt(content)
            elif file_extension == ".pdf":
                return await self._extract_from_pdf(content)
            elif file_extension == ".docx":
                return await self._extract_from_docx(content)
            else:
                raise DocumentProcessingError(
                    f"Unsupported file type for text extraction: {file_extension}",
                    filename=filename,
                    error_type="unsupported_extraction"
                )
                
        except DocumentProcessingError as e:
            # Try fallback methods if primary extraction fails
            logger.warning(
                "primary_extraction_failed_trying_fallback",
                filename=filename,
                error=str(e)
            )
            
            # Attempt fallback extraction
            fallback_result = await self._fallback_extraction(filename, content, file_extension, str(e))
            if fallback_result[0]:  # If some text was extracted
                return fallback_result
            
            # If fallback also failed, return error
            error_msg = f"Text extraction failed: {str(e)}"
            return "", error_msg
            
        except Exception as e:
            error_msg = f"Text extraction failed: {str(e)}"
            logger.error(
                "text_extraction_failed",
                filename=filename,
                file_type=file_extension,
                error=str(e),
                exc_info=True
            )
            return "", error_msg
    
    async def _fallback_extraction(
        self,
        filename: str,
        content: bytes,
        file_extension: str,
        original_error: str
    ) -> Tuple[str, Optional[str]]:
        """
        Fallback extraction methods when primary extraction fails.
        
        Returns:
            Tuple of (extracted_text, warning_message)
        """
        extracted_text = ""
        warnings = [f"Primary extraction failed: {original_error}"]
        
        try:
            if file_extension == ".pdf":
                # Try alternative PDF extraction
                extracted_text = await self._fallback_pdf_extraction(content)
                warnings.append("Used fallback PDF extraction (basic text extraction)")
                
            elif file_extension == ".docx":
                # Try extracting as plain ZIP
                extracted_text = await self._fallback_docx_extraction(content)
                warnings.append("Used fallback DOCX extraction (raw XML parsing)")
                
            elif file_extension == ".txt":
                # Try with more aggressive encoding detection
                extracted_text = await self._fallback_txt_extraction(content)
                warnings.append("Used fallback text extraction (forced UTF-8 with replacements)")
                
        except Exception as e:
            warnings.append(f"Fallback extraction also failed: {str(e)}")
        
        warning_msg = "; ".join(warnings) if warnings else None
        return extracted_text, warning_msg
    
    async def _fallback_pdf_extraction(self, content: bytes) -> str:
        """Fallback PDF extraction using basic text search."""
        try:
            # Try to extract readable text using regex
            import re
            
            # Decode with error handling
            text = content.decode('latin-1', errors='ignore')
            
            # Extract text between common PDF markers
            text_pattern = re.compile(r'BT\s*(.*?)\s*ET', re.DOTALL)
            matches = text_pattern.findall(text)
            
            extracted_parts = []
            for match in matches:
                # Clean up PDF commands
                cleaned = re.sub(r'[<>\[\]()]+', ' ', match)
                cleaned = re.sub(r'\d+\s+\d+\s+Td', ' ', cleaned)
                cleaned = re.sub(r'Tj', ' ', cleaned)
                cleaned = re.sub(r'\s+', ' ', cleaned)
                
                if len(cleaned.strip()) > 5:  # Minimum meaningful text
                    extracted_parts.append(cleaned.strip())
            
            return ' '.join(extracted_parts)
            
        except Exception as e:
            logger.error("fallback_pdf_extraction_failed", error=str(e))
            return ""
    
    async def _fallback_docx_extraction(self, content: bytes) -> str:
        """Fallback DOCX extraction by parsing raw XML."""
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            import io
            
            # Open as ZIP
            zip_file = io.BytesIO(content)
            extracted_text = []
            
            with zipfile.ZipFile(zip_file, 'r') as docx:
                # Try to read document.xml directly
                if 'word/document.xml' in docx.namelist():
                    xml_content = docx.read('word/document.xml')
                    
                    # Parse XML and extract text
                    root = ET.fromstring(xml_content)
                    
                    # Find all text elements (w:t tags)
                    for elem in root.iter():
                        if elem.tag.endswith('}t'):  # w:t tags
                            if elem.text:
                                extracted_text.append(elem.text)
            
            return ' '.join(extracted_text)
            
        except Exception as e:
            logger.error("fallback_docx_extraction_failed", error=str(e))
            return ""
    
    async def _fallback_txt_extraction(self, content: bytes) -> str:
        """Fallback text extraction with aggressive decoding."""
        try:
            # Try multiple encodings in order of likelihood
            encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16', 'ascii']
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    # If successful, return
                    return text
                except:
                    continue
            
            # Last resort: decode with replacements
            return content.decode('utf-8', errors='replace')
            
        except Exception as e:
            logger.error("fallback_txt_extraction_failed", error=str(e))
            return ""
    
    async def _extract_from_txt(self, content: bytes) -> Tuple[str, None]:
        """Extract text from TXT file with encoding validation."""
        try:
            # Check for binary content (not text)
            if b'\x00' in content[:1024]:  # Null bytes in first 1KB
                raise DocumentProcessingError("File appears to be binary, not text")
            
            # Try UTF-8 first, then fall back to other encodings
            text = None
            encoding_used = None
            
            try:
                text = content.decode('utf-8')
                encoding_used = 'utf-8'
            except UnicodeDecodeError:
                # Try to detect encoding
                import chardet
                detected = chardet.detect(content)
                encoding = detected.get('encoding')
                confidence = detected.get('confidence', 0)
                
                if not encoding or confidence < 0.7:
                    # Low confidence in encoding detection
                    raise DocumentProcessingError(
                        f"Unable to detect text encoding (confidence: {confidence:.2%})"
                    )
                
                try:
                    text = content.decode(encoding)
                    encoding_used = encoding
                except Exception as e:
                    # Last resort: try with error replacement
                    text = content.decode(encoding, errors='replace')
                    encoding_used = f"{encoding} (with replacements)"
            
            # Validate extracted text
            if not text or len(text.strip()) == 0:
                raise DocumentProcessingError("No text content found in file")
            
            # Normalize line endings
            text = text.replace('\r\n', '\n').replace('\r', '\n')
            
            logger.info("txt_extraction_successful", encoding=encoding_used, chars=len(text))
            return text.strip(), None
            
        except DocumentProcessingError:
            raise
        except Exception as e:
            return "", f"TXT extraction failed: {str(e)}"
    
    async def _extract_from_pdf(self, content: bytes) -> Tuple[str, None]:
        """Extract text from PDF file with enhanced corruption detection and memory monitoring."""
        try:
            import PyPDF2
            import io
            
            # Check memory before starting
            self.check_memory_limit("PDF extraction")
            
            # Pre-extraction validation
            if len(content) < 100:  # Minimum viable PDF size
                raise DocumentProcessingError("PDF file too small to be valid")
            
            # Verify PDF structure
            if not content.startswith(b'%PDF'):
                raise DocumentProcessingError("Invalid PDF: missing PDF header")
            
            # Check for EOF marker (should be near the end)
            eof_found = False
            for marker in [b'%%EOF', b'%%EOF\n', b'%%EOF\r\n', b'%%EOF\r']:
                if marker in content[-1024:]:
                    eof_found = True
                    break
            
            if not eof_found:
                logger.warning("pdf_missing_eof_marker", size=len(content))
                # Don't fail, but log warning - some PDFs may still work
            
            pdf_file = io.BytesIO(content)
            
            # Attempt to create reader with error handling
            try:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
            except Exception as e:
                raise DocumentProcessingError(f"Corrupted PDF structure: {str(e)}")
            
            # Validate PDF properties
            num_pages = len(pdf_reader.pages)
            if num_pages == 0:
                raise DocumentProcessingError("PDF contains no pages")
            
            # Check if encrypted
            if pdf_reader.is_encrypted:
                try:
                    # Try empty password
                    if not pdf_reader.decrypt(''):
                        raise DocumentProcessingError("PDF is password-protected")
                except:
                    raise DocumentProcessingError("Cannot decrypt PDF file")
            
            # Extract text with quality tracking and partial extraction support
            text_parts = []
            failed_pages = []
            total_chars = 0
            max_pages = getattr(settings, 'MAX_PDF_PAGES', 1000)  # Configurable limit
            partial_extraction = False
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    # Check memory periodically
                    if page_num % 10 == 0:  # Every 10 pages
                        self.check_memory_limit(f"PDF extraction (page {page_num + 1})")
                    
                    # Check if we've hit page limit
                    if page_num >= max_pages:
                        logger.warning(
                            "pdf_page_limit_reached",
                            extracted_pages=page_num,
                            total_pages=num_pages,
                            limit=max_pages
                        )
                        partial_extraction = True
                        break
                    
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text)
                        total_chars += len(page_text)
                        
                        # Check if we have enough content (partial extraction)
                        if total_chars > 1_000_000:  # 1MB of text
                            logger.info(
                                "partial_extraction_size_limit",
                                pages_extracted=page_num + 1,
                                total_chars=total_chars
                            )
                            partial_extraction = True
                            break
                            
                except DocumentProcessingError:
                    raise  # Re-raise memory errors
                except Exception as e:
                    failed_pages.append(page_num + 1)
                    logger.warning(
                        "pdf_page_extraction_failed",
                        page_number=page_num + 1,
                        error=str(e)
                    )
                    # Continue with partial extraction
                    continue
            
            # Validate extraction quality
            if total_chars == 0:
                raise DocumentProcessingError(
                    f"No text extracted from {num_pages} page(s). PDF may be image-based or corrupted."
                )
            
            # Build result
            extracted_text = "\n\n".join(text_parts)
            
            # Build warning message
            warnings = []
            if failed_pages:
                warnings.append(f"Failed to extract {len(failed_pages)} page(s): {failed_pages[:5]}{'...' if len(failed_pages) > 5 else ''}")
            
            if partial_extraction:
                extracted_pages = len(text_parts)
                warnings.append(f"Partial extraction: {extracted_pages} of {num_pages} pages processed")
            
            warning = "; ".join(warnings) if warnings else None
            
            logger.info(
                "pdf_extraction_complete",
                pages=num_pages,
                extracted_pages=len(text_parts),
                total_chars=total_chars,
                failed_pages=len(failed_pages)
            )
            
            return extracted_text.strip(), warning
            
        except DocumentProcessingError:
            raise
        except Exception as e:
            return "", f"PDF extraction failed: {str(e)}"
    
    async def _extract_from_docx(self, content: bytes) -> Tuple[str, None]:
        """Extract text from DOCX file with corruption detection and memory monitoring."""
        try:
            import docx
            import io
            import zipfile
            
            # Check memory before starting
            self.check_memory_limit("DOCX extraction")
            
            # Pre-extraction validation
            if len(content) < 100:  # Minimum viable DOCX size
                raise DocumentProcessingError("DOCX file too small to be valid")
            
            # DOCX files are ZIP archives
            if not content.startswith(b'PK'):
                raise DocumentProcessingError("Invalid DOCX: not a valid ZIP archive")
            
            # Verify it's a valid ZIP
            try:
                doc_file = io.BytesIO(content)
                # Test if we can read it as ZIP
                with zipfile.ZipFile(doc_file, 'r') as zip_test:
                    # Check for required DOCX structure
                    required_files = ['word/document.xml', '[Content_Types].xml']
                    zip_contents = zip_test.namelist()
                    
                    for required in required_files:
                        if required not in zip_contents:
                            raise DocumentProcessingError(
                                f"Invalid DOCX structure: missing {required}"
                            )
                
                # Reset file pointer
                doc_file.seek(0)
            except zipfile.BadZipFile:
                raise DocumentProcessingError("Corrupted DOCX: invalid ZIP structure")
            except Exception as e:
                raise DocumentProcessingError(f"DOCX validation failed: {str(e)}")
            
            # Attempt to parse as DOCX
            try:
                doc = docx.Document(doc_file)
            except Exception as e:
                raise DocumentProcessingError(f"Corrupted DOCX document: {str(e)}")
            
            # Extract text with quality tracking
            text_parts = []
            total_chars = 0
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
                    total_chars += len(text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_row = " | ".join(row_text)
                        text_parts.append(table_row)
                        total_chars += len(table_row)
            
            # Validate extraction
            if total_chars == 0:
                raise DocumentProcessingError("No text content found in DOCX file")
            
            extracted_text = "\n\n".join(text_parts)
            
            logger.info(
                "docx_extraction_complete",
                paragraphs=len(doc.paragraphs),
                tables=len(doc.tables),
                total_chars=total_chars
            )
            
            return extracted_text.strip(), None
            
        except DocumentProcessingError:
            raise
        except Exception as e:
            return "", f"DOCX extraction failed: {str(e)}"
    
    async def create_document_record(
        self,
        session: AsyncSession,
        filename: str,
        original_filename: str,
        content: bytes,
        document_data: Optional[DocumentCreate] = None
    ) -> Document:
        """
        Create document record in database.
        
        Args:
            session: Database session
            filename: Safe filename for storage
            original_filename: Original uploaded filename
            content: File content
            document_data: Optional additional document data
            
        Returns:
            Created Document instance
        """
        file_hash = self.calculate_file_hash(content)
        file_extension = Path(original_filename).suffix.lower()
        file_size = len(content)
        
        # Check for duplicate by hash
        existing_doc = await session.execute(
            select(Document).where(Document.file_hash == file_hash)
        )
        if existing_doc.scalar_one_or_none():
            raise DocumentProcessingError(
                f"Document with identical content already exists",
                filename=original_filename,
                error_type="duplicate_content"
            )
        
        # Create document record
        document = Document(
            filename=filename,
            original_filename=original_filename,
            file_type=file_extension,
            file_size_bytes=file_size,
            file_hash=file_hash,
            processing_status=ProcessingStatus.PENDING.value,
            total_chunks=0
        )
        
        # Add optional data if provided
        if document_data:
            if document_data.tags:
                document.tags = document_data.tags
            if document_data.document_metadata:
                document.document_metadata = document_data.document_metadata
        
        session.add(document)
        await session.flush()  # Get the ID
        
        logger.info(
            "document_record_created",
            document_id=str(document.id),
            filename=filename,
            original_filename=original_filename,
            file_size_bytes=file_size,
            file_hash=file_hash
        )
        
        return document
    
    async def update_processing_status(
        self,
        session: AsyncSession,
        document_id: str,
        status: ProcessingStatus,
        error_message: Optional[str] = None
    ) -> None:
        """Update document processing status."""
        update_data = {"processing_status": status.value}
        
        if status == ProcessingStatus.PROCESSING:
            # REFACTORED: Using existing utility instead of direct datetime.utcnow()
            update_data["processing_started_at"] = get_utc_datetime()
        elif status in [ProcessingStatus.COMPLETED, ProcessingStatus.FAILED]:
            # REFACTORED: Using existing utility instead of direct datetime.utcnow()
            update_data["processing_completed_at"] = get_utc_datetime()
            
        if error_message:
            update_data["processing_error"] = error_message
        
        await session.execute(
            update(Document)
            .where(Document.id == document_id)
            .values(**update_data)
        )
        
        logger.info(
            "document_status_updated",
            document_id=document_id,
            status=status.value,
            error=error_message
        )
    
    @asynccontextmanager
    async def processing_transaction(
        self,
        session: AsyncSession,
        document_id: Optional[str] = None
    ):
        """
        Context manager for atomic document processing operations.
        
        Usage:
            async with processor.processing_transaction(session, doc_id) as transaction:
                # Do processing operations
                transaction.add_file(filepath)
                transaction.add_chunk(chunk_id)
                # If exception occurs, automatic rollback
        """
        transaction = ProcessingTransaction(
            document_id=document_id,
            session=session
        )
        
        try:
            yield transaction
            # If we get here, processing succeeded
            transaction.commit()
        except Exception as e:
            # Rollback on any error
            logger.error(
                "processing_transaction_failed",
                document_id=document_id,
                error=str(e),
                exc_info=True
            )
            
            rollback_success = await transaction.rollback()
            if not rollback_success:
                logger.error(
                    "transaction_rollback_had_errors",
                    errors=transaction.rollback_errors
                )
            
            # Re-raise the original exception
            raise
    
    async def save_file_with_transaction(
        self,
        content: bytes,
        filename: str,
        transaction: ProcessingTransaction,
        storage_path: str = "/tmp"
    ) -> str:
        """
        Save file and track it in the transaction.
        
        Returns:
            Full path to saved file
        """
        filepath = os.path.join(storage_path, filename)
        
        try:
            # Ensure directory exists
            os.makedirs(storage_path, exist_ok=True)
            
            # Write file
            with open(filepath, 'wb') as f:
                f.write(content)
            
            # Track in transaction
            transaction.add_file(filepath)
            
            logger.info(
                "file_saved_in_transaction",
                filepath=filepath,
                size_bytes=len(content)
            )
            
            return filepath
            
        except Exception as e:
            logger.error(
                "file_save_failed",
                filepath=filepath,
                error=str(e)
            )
            raise DocumentProcessingError(f"Failed to save file: {str(e)}")
    
    async def process_document_async(self, document_id: str) -> None:
        """
        Asynchronously process a document with full transaction support and metrics.
        """
        # Start metrics tracking
        metrics = metrics_service.start_processing(document_id)
        
        try:
            async with db_manager.get_session() as session:
                # Create processing transaction
                async with self.processing_transaction(session, document_id) as transaction:
                    try:
                        # Get document
                        result = await session.execute(
                            select(Document).where(Document.id == document_id)
                        )
                        document = result.scalar_one_or_none()
                        if not document:
                            logger.error("document_not_found", document_id=document_id)
                            return
                        
                        # Update file size in metrics
                        if document.file_size_bytes:
                            metrics.file_size_bytes = document.file_size_bytes
                        
                        # Update status to processing
                        await self.update_processing_status(
                            session, document_id, ProcessingStatus.PROCESSING
                        )
                        await session.flush()  # Don't commit yet
                        
                        # TODO: Full processing pipeline with transaction and metrics tracking:
                        
                        # 1. Load file content from storage
                        with metrics_service.track_stage(document_id, "storage"):
                            # Simulate loading
                            await asyncio.sleep(0.1)
                            # Record memory sample
                            current_memory = self.get_memory_usage_mb()
                            metrics_service.record_memory_sample(document_id, current_memory)
                        
                        # 2. Extract text content
                        with metrics_service.track_stage(document_id, "extraction"):
                            # Simulate extraction
                            await asyncio.sleep(0.2)
                            # Record extraction metrics
                            metrics_service.record_processing_stats(
                                document_id,
                                extracted_chars=10000  # Simulated
                            )
                            metrics_service.record_quality_metrics(
                                document_id,
                                extraction_quality=0.95  # Simulated
                            )
                        
                        # 3. Create chunks using chunking service
                        with metrics_service.track_stage(document_id, "chunking"):
                            # Simulate chunking
                            await asyncio.sleep(0.1)
                            # Record chunk metrics
                            metrics_service.record_processing_stats(
                                document_id,
                                total_chunks=50,  # Simulated
                                quality_chunks=45,
                                duplicate_chunks=5
                            )
                            # Track chunk IDs with transaction
                            for i in range(5):  # Simulated chunk IDs
                                transaction.add_chunk(f"chunk_{i}")
                        
                        # 4. Create embeddings using vector store service
                        with metrics_service.track_stage(document_id, "embedding"):
                            # Simulate embedding
                            await asyncio.sleep(0.3)
                            # Track embeddings with transaction
                            for i in range(5):  # Simulated embedding IDs
                                transaction.add_embedding(f"embedding_{i}")
                        
                        # 5. Store chunks in database (already tracked above)
                        
                        # Update to completed
                        await self.update_processing_status(
                            session, document_id, ProcessingStatus.COMPLETED
                        )
                        
                        # Commit the database transaction
                        await session.commit()
                        
                        logger.info(
                            "document_processing_completed",
                            document_id=document_id
                        )
                        
                    except Exception as e:
                        # Transaction will automatically rollback
                        logger.error(
                            "document_processing_failed",
                            document_id=document_id,
                            error=str(e),
                            exc_info=True
                        )
                        
                        # Update status to failed
                        await self.update_processing_status(
                            session, document_id, ProcessingStatus.FAILED, str(e)
                        )
                        await session.commit()
                        
                        raise  # Re-raise to trigger transaction rollback
        
        finally:
            # Always end metrics tracking
            final_metrics = metrics_service.end_processing(document_id)
            
            if final_metrics:
                logger.info(
                    "processing_metrics_summary",
                    document_id=document_id,
                    success=final_metrics.overall_success,
                    total_time_seconds=final_metrics.total_time,
                    peak_memory_mb=final_metrics.peak_memory_mb
                )


# Global processor instance
document_processor = DocumentProcessor()