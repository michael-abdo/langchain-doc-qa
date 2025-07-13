"""
Document processing service.
Handles file upload, validation, content extraction, and processing orchestration.
"""
import os
import hashlib
import mimetypes
from typing import Optional, Tuple, Dict, Any, List
from pathlib import Path
import asyncio
from datetime import datetime

from app.core.config import settings
from app.core.logging import get_logger, get_utc_datetime
from app.core.database import db_manager
from app.core.exceptions import (
    DocumentProcessingError,
    FileTooLargeError,
    UnsupportedFileTypeError,
    ValidationError
)
from app.models.document import Document, DocumentChunk
from app.schemas.document import DocumentCreate, ProcessingStatus
from sqlalchemy import select, update
from sqlalchemy.ext.asyncio import AsyncSession

logger = get_logger(__name__)


class DocumentProcessor:
    """Handles document processing operations."""
    
    def __init__(self):
        self.supported_types = set(settings.ALLOWED_FILE_TYPES)
        self.max_size_bytes = settings.MAX_FILE_SIZE_MB * 1024 * 1024
        
    async def validate_file(self, filename: str, content: bytes) -> None:
        """
        Validate uploaded file meets requirements.
        
        Args:
            filename: Original filename
            content: File content as bytes
            
        Raises:
            FileTooLargeError: If file exceeds size limit
            UnsupportedFileTypeError: If file type not supported
        """
        # Check file size
        file_size_bytes = len(content)
        if file_size_bytes > self.max_size_bytes:
            file_size_mb = file_size_bytes / (1024 * 1024)
            raise FileTooLargeError(
                filename=filename,
                size_mb=file_size_mb,
                max_size_mb=settings.MAX_FILE_SIZE_MB
            )
        
        # Check file extension
        file_extension = Path(filename).suffix.lower()
        if file_extension not in self.supported_types:
            raise UnsupportedFileTypeError(
                filename=filename,
                file_type=file_extension,
                supported_types=list(self.supported_types)
            )
        
        logger.info(
            "file_validation_passed",
            filename=filename,
            size_bytes=file_size_bytes,
            file_type=file_extension
        )
    
    def calculate_file_hash(self, content: bytes) -> str:
        """Calculate SHA-256 hash of file content."""
        return hashlib.sha256(content).hexdigest()
    
    def generate_safe_filename(self, original_filename: str) -> str:
        """Generate a safe filename for storage."""
        # Remove unsafe characters and create unique filename
        safe_name = "".join(c for c in original_filename if c.isalnum() or c in "._-")
        # REFACTORED: Using existing utility instead of direct datetime.utcnow()
        timestamp = get_utc_datetime().strftime("%Y%m%d_%H%M%S")
        name_part, ext_part = os.path.splitext(safe_name)
        return f"{timestamp}_{name_part[:50]}{ext_part}"
    
    async def extract_text_content(self, filename: str, content: bytes) -> Tuple[str, Optional[str]]:
        """
        Extract text content from file.
        
        Args:
            filename: Original filename
            content: File content as bytes
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        file_extension = Path(filename).suffix.lower()
        
        try:
            if file_extension == ".txt":
                return await self._extract_from_txt(content)
            elif file_extension == ".pdf":
                return await self._extract_from_pdf(content)
            elif file_extension == ".docx":
                return await self._extract_from_docx(content)
            else:
                raise DocumentProcessingError(
                    f"Unsupported file type for text extraction: {file_extension}",
                    filename=filename,
                    error_type="unsupported_extraction"
                )
                
        except Exception as e:
            error_msg = f"Text extraction failed: {str(e)}"
            logger.error(
                "text_extraction_failed",
                filename=filename,
                file_type=file_extension,
                error=str(e),
                exc_info=True
            )
            return "", error_msg
    
    async def _extract_from_txt(self, content: bytes) -> Tuple[str, None]:
        """Extract text from TXT file."""
        try:
            # Try UTF-8 first, then fall back to other encodings
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                # Try to detect encoding
                import chardet
                detected = chardet.detect(content)
                encoding = detected.get('encoding', 'utf-8')
                text = content.decode(encoding, errors='replace')
            
            return text.strip(), None
            
        except Exception as e:
            return "", f"TXT extraction failed: {str(e)}"
    
    async def _extract_from_pdf(self, content: bytes) -> Tuple[str, None]:
        """Extract text from PDF file."""
        try:
            import PyPDF2
            import io
            
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(
                        "pdf_page_extraction_failed",
                        page_number=page_num,
                        error=str(e)
                    )
                    continue
            
            extracted_text = "\n\n".join(text_parts)
            return extracted_text.strip(), None
            
        except Exception as e:
            return "", f"PDF extraction failed: {str(e)}"
    
    async def _extract_from_docx(self, content: bytes) -> Tuple[str, None]:
        """Extract text from DOCX file."""
        try:
            import docx
            import io
            
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            extracted_text = "\n\n".join(text_parts)
            return extracted_text.strip(), None
            
        except Exception as e:
            return "", f"DOCX extraction failed: {str(e)}"
    
    async def create_document_record(
        self,
        session: AsyncSession,
        filename: str,
        original_filename: str,
        content: bytes,
        document_data: Optional[DocumentCreate] = None
    ) -> Document:
        """
        Create document record in database.
        
        Args:
            session: Database session
            filename: Safe filename for storage
            original_filename: Original uploaded filename
            content: File content
            document_data: Optional additional document data
            
        Returns:
            Created Document instance
        """
        file_hash = self.calculate_file_hash(content)
        file_extension = Path(original_filename).suffix.lower()
        file_size = len(content)
        
        # Check for duplicate by hash
        existing_doc = await session.execute(
            select(Document).where(Document.file_hash == file_hash)
        )
        if existing_doc.scalar_one_or_none():
            raise DocumentProcessingError(
                f"Document with identical content already exists",
                filename=original_filename,
                error_type="duplicate_content"
            )
        
        # Create document record
        document = Document(
            filename=filename,
            original_filename=original_filename,
            file_type=file_extension,
            file_size_bytes=file_size,
            file_hash=file_hash,
            processing_status=ProcessingStatus.PENDING.value,
            total_chunks=0
        )
        
        # Add optional data if provided
        if document_data:
            if document_data.tags:
                document.tags = document_data.tags
            if document_data.document_metadata:
                document.document_metadata = document_data.document_metadata
        
        session.add(document)
        await session.flush()  # Get the ID
        
        logger.info(
            "document_record_created",
            document_id=str(document.id),
            filename=filename,
            original_filename=original_filename,
            file_size_bytes=file_size,
            file_hash=file_hash
        )
        
        return document
    
    async def update_processing_status(
        self,
        session: AsyncSession,
        document_id: str,
        status: ProcessingStatus,
        error_message: Optional[str] = None
    ) -> None:
        """Update document processing status."""
        update_data = {"processing_status": status.value}
        
        if status == ProcessingStatus.PROCESSING:
            # REFACTORED: Using existing utility instead of direct datetime.utcnow()
            update_data["processing_started_at"] = get_utc_datetime()
        elif status in [ProcessingStatus.COMPLETED, ProcessingStatus.FAILED]:
            # REFACTORED: Using existing utility instead of direct datetime.utcnow()
            update_data["processing_completed_at"] = get_utc_datetime()
            
        if error_message:
            update_data["processing_error"] = error_message
        
        await session.execute(
            update(Document)
            .where(Document.id == document_id)
            .values(**update_data)
        )
        
        logger.info(
            "document_status_updated",
            document_id=document_id,
            status=status.value,
            error=error_message
        )
    
    async def process_document_async(self, document_id: str) -> None:
        """
        Asynchronously process a document.
        This would typically be called by a background task queue.
        """
        async with db_manager.get_session() as session:
            try:
                # Get document
                result = await session.execute(
                    select(Document).where(Document.id == document_id)
                )
                document = result.scalar_one_or_none()
                if not document:
                    logger.error("document_not_found", document_id=document_id)
                    return
                
                # Update status to processing
                await self.update_processing_status(
                    session, document_id, ProcessingStatus.PROCESSING
                )
                await session.commit()
                
                # TODO: This is where we would:
                # 1. Load file content from storage
                # 2. Extract text content
                # 3. Create chunks using chunking service
                # 4. Create embeddings using vector store service
                # 5. Store chunks in database
                
                # For now, simulate processing
                await asyncio.sleep(1)
                
                # Update to completed (this would be done after actual processing)
                await self.update_processing_status(
                    session, document_id, ProcessingStatus.COMPLETED
                )
                await session.commit()
                
                logger.info(
                    "document_processing_completed",
                    document_id=document_id
                )
                
            except Exception as e:
                logger.error(
                    "document_processing_failed",
                    document_id=document_id,
                    error=str(e),
                    exc_info=True
                )
                
                # Update status to failed
                await self.update_processing_status(
                    session, document_id, ProcessingStatus.FAILED, str(e)
                )
                await session.commit()


# Global processor instance
document_processor = DocumentProcessor()